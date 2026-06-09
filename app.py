# =========================================
# ENGLISH REPORT COMMENT GENERATOR - Streamlit Version
# =========================================

import random
import streamlit as st
from docx import Document

TARGET_CHARS = 499  # target character count including spaces

from statements import (
    opening_phrases,
    attitude_bank,
    reading_bank,
    writing_bank,
    reading_target_bank,
    writing_target_bank,
    closer_bank
)

from eoy_statements import (
    eoy_opening_phrases,
    eoy_attitude_bank,
    eoy_reading_bank,
    eoy_writing_bank,
    eoy_reading_target_bank,
    eoy_writing_target_bank,
    eoy_closer_bank
)


# ---------- HELPERS ----------
def get_pronouns(gender):
    gender = gender.lower()
    if gender == "male":
        return "he", "his"
    elif gender == "female":
        return "she", "her"
    return "they", "their"

def lowercase_first(text):
    return text[0].lower() + text[1:] if text else ""

def truncate_comment(comment, target=TARGET_CHARS):
    if len(comment) <= target:
        return comment
    truncated = comment[:target].rstrip(" ,;.")
    if "." in truncated:
        truncated = truncated[:truncated.rfind(".")+1]
    return truncated

def replace_name_placeholder(text, name):
    return text.replace("{name}", name)

def generate_comment(name, att, read, write, read_t, write_t, pronouns, attitude_target=None):
    p, p_poss = pronouns
    opening = random.choice(opening_phrases)

    attitude_sentence = f"{opening} {name} {attitude_bank[att]}."
    reading_sentence = f"In reading, {p} {reading_bank[read]}."
    writing_sentence = f"In writing, {p} {writing_bank[write]}."
    reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_bank[read_t])}."
    writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_bank[write_t])}."
    closer_sentence = random.choice(closer_bank)

    # FIX: attitude target now placed after writing target, before closer
    comment_parts = [
        attitude_sentence,
        reading_sentence,
        writing_sentence,
        reading_target_sentence,
        writing_target_sentence,
    ]

    if attitude_target:
        attitude_target_sentence = attitude_target.strip()
        if not attitude_target_sentence.endswith("."):
            attitude_target_sentence += "."
        comment_parts.append(attitude_target_sentence)

    comment_parts.append(closer_sentence)

    comment = " ".join(comment_parts)
    comment = truncate_comment(comment, TARGET_CHARS)
    return comment

def generate_eoy_comment(name, att, read, write, read_t, write_t, pronouns, attitude_target=None):
    p, p_poss = pronouns
    opening = random.choice(eoy_opening_phrases)

    # Gender-correct the attitude bank statements (they use "her" as placeholder)
    # Replace "her" with correct possessive and "she" with correct pronoun
    def fix_gender(text):
        if p == "he":
            text = text.replace(" she ", " he ").replace(" her ", " his ").replace(" her.", " him.")
            text = text.replace("her own", "his own").replace("her progress", "his progress")
            text = text.replace("her engagement", "his engagement").replace("her confidence", "his confidence")
            text = text.replace("her ability", "his ability").replace("her approach", "his approach")
        elif p == "they":
            text = text.replace(" she ", " they ").replace(" her ", " their ").replace(" her.", " them.")
            text = text.replace("her own", "their own").replace("her progress", "their progress")
            text = text.replace("her engagement", "their engagement").replace("her confidence", "their confidence")
            text = text.replace("her ability", "their ability").replace("her approach", "their approach")
        return text

    attitude_sentence = f"{opening} {name} {fix_gender(eoy_attitude_bank[att])}."
    reading_sentence = f"In reading, {p} {fix_gender(eoy_reading_bank[read])}."
    writing_sentence = f"In writing, {p} {fix_gender(eoy_writing_bank[write])}."
    reading_target_sentence = f"Going into Year 8, {p} should {lowercase_first(eoy_reading_target_bank[read_t])}."
    writing_target_sentence = f"In writing, {p} should {lowercase_first(eoy_writing_target_bank[write_t])}."

    closer_raw = random.choice(eoy_closer_bank)
    closer_sentence = replace_name_placeholder(closer_raw, name)

    comment_parts = [
        attitude_sentence,
        reading_sentence,
        writing_sentence,
        reading_target_sentence,
        writing_target_sentence,
    ]

    if attitude_target:
        attitude_target_sentence = attitude_target.strip()
        if not attitude_target_sentence.endswith("."):
            attitude_target_sentence += "."
        comment_parts.append(attitude_target_sentence)

    comment_parts.append(closer_sentence)

    comment = " ".join(comment_parts)
    comment = truncate_comment(comment, TARGET_CHARS)
    return comment


# ---------- STREAMLIT APP ----------
st.title("English Report Comment Generator (~499 chars)")

tab1, tab2 = st.tabs(["Termly Comment", "End of Year Comment"])

EOY_BANDS = [100, 90, 80, 70, 60, 50, 40, 30]

# ---------- TAB 1: TERMLY ----------
with tab1:
    st.markdown("Fill in the student details and click **Generate Comment**.")

    if 'all_comments' not in st.session_state:
        st.session_state['all_comments'] = []

    with st.form("report_form"):
        name = st.text_input("Student Name")
        gender = st.selectbox("Gender", ["Male", "Female"])
        att = st.selectbox("Attitude band", [90, 85, 80, 75, 70, 65, 60, 55, 40])
        read = st.selectbox("Reading achievement band", [90, 85, 80, 75, 70, 65, 60, 55, 40])
        write = st.selectbox("Writing achievement band", [90, 85, 80, 75, 70, 65, 60, 55, 40])
        read_t = st.selectbox("Reading target band", [90, 85, 80, 75, 70, 65, 60, 55, 40])
        write_t = st.selectbox("Writing target band", [90, 85, 80, 75, 70, 65, 60, 55, 40])
        attitude_target = st.text_input("Optional Attitude Next Steps (appears last)")
        submitted = st.form_submit_button("Generate Comment")

    if submitted and name:
        pronouns = get_pronouns(gender)
        comment = generate_comment(name, att, read, write, read_t, write_t, pronouns, attitude_target)
        char_count = len(comment)
        st.text_area("Generated Comment", comment, height=200)
        st.write(f"Character count (including spaces): {char_count} / {TARGET_CHARS}")
        st.session_state['all_comments'].append(f"{name}: {comment}")

    if st.session_state.get('all_comments'):
        st.markdown("### All Generated Comments:")
        for c in st.session_state['all_comments']:
            st.write(c)

        if st.button("Download Termly Report (Word)"):
            doc = Document()
            for c in st.session_state['all_comments']:
                doc.add_paragraph(c)
            file_name = "English_Report_Comments.docx"
            doc.save(file_name)
            with open(file_name, "rb") as f:
                st.download_button(
                    label="Download Word File",
                    data=f,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


# ---------- TAB 2: END OF YEAR ----------
with tab2:
    st.markdown("End of Year comments for Year 7. Warmer, reflective tone with growth-focused language.")

    if 'eoy_comments' not in st.session_state:
        st.session_state['eoy_comments'] = []

    with st.form("eoy_form"):
        eoy_name = st.text_input("Student Name", key="eoy_name")
        eoy_gender = st.selectbox("Gender", ["Male", "Female"], key="eoy_gender")
        eoy_att = st.selectbox("Attitude band", EOY_BANDS, key="eoy_att")
        eoy_read = st.selectbox("Reading achievement band", EOY_BANDS, key="eoy_read")
        eoy_write = st.selectbox("Writing achievement band", EOY_BANDS, key="eoy_write")
        eoy_read_t = st.selectbox("Reading target band", EOY_BANDS, key="eoy_read_t")
        eoy_write_t = st.selectbox("Writing target band", EOY_BANDS, key="eoy_write_t")
        eoy_attitude_target = st.text_input("Optional Attitude Next Steps (appears last)", key="eoy_att_target")
        eoy_submitted = st.form_submit_button("Generate End of Year Comment")

    if eoy_submitted and eoy_name:
        pronouns = get_pronouns(eoy_gender)
        comment = generate_eoy_comment(eoy_name, eoy_att, eoy_read, eoy_write, eoy_read_t, eoy_write_t, pronouns, eoy_attitude_target)
        char_count = len(comment)
        st.text_area("Generated Comment", comment, height=200)
        st.write(f"Character count (including spaces): {char_count} / {TARGET_CHARS}")
        st.session_state['eoy_comments'].append(f"{eoy_name}: {comment}")

    if st.session_state.get('eoy_comments'):
        st.markdown("### All Generated End of Year Comments:")
        for c in st.session_state['eoy_comments']:
            st.write(c)

        if st.button("Download End of Year Report (Word)"):
            doc = Document()
            for c in st.session_state['eoy_comments']:
                doc.add_paragraph(c)
            file_name = "EOY_English_Report_Comments.docx"
            doc.save(file_name)
            with open(file_name, "rb") as f:
                st.download_button(
                    label="Download Word File",
                    data=f,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
